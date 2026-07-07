import json
import subprocess
import sys

from docx import Document

from tools import read_docx as rd


def _build_docx(path):
    doc = Document()
    doc.add_heading("Заголовок ролика", level=0)
    doc.add_paragraph("Публикация: 2026-06-25 14:00 MSK")
    doc.add_paragraph("Первый абзац")
    doc.add_paragraph("ВИЗУАЛ НА ЭКРАН: таблица")
    doc.add_paragraph("СХЕМА НА ЭКРАН: процесс")
    doc.add_paragraph("КАРТИНКА НА ЭКРАН: график")
    doc.add_paragraph("— — —")
    doc.add_paragraph("Второй абзац")
    doc.save(path)


def test_parse_filename_supports_russian_title_schedule():
    scheduled_at, alias = rd.parse_filename("2026-06-25-1400-Топ 5 экономических инструментов.docx")

    assert scheduled_at == "2026-06-25 14:00"
    assert alias == "Топ 5 экономических инструментов"


def test_parse_filename_supports_legacy_date_only_name():
    scheduled_at, alias = rd.parse_filename("2026-06-25-Старый формат.docx")

    assert scheduled_at is None
    assert alias == "Старый формат"


def test_read_docx_filters_metadata_and_visual_markers(tmp_path):
    docx_path = tmp_path / "2026-06-25-1400-Тест.docx"
    _build_docx(docx_path)

    title, body = rd.read_docx(docx_path)

    assert title == "Заголовок ролика"
    assert body == "Первый абзац\n\nВторой абзац"
    assert "Публикация:" not in body
    assert "ВИЗУАЛ НА ЭКРАН" not in body


def test_read_docx_cli_outputs_expected_json(tmp_path):
    docx_path = tmp_path / "2026-06-25-1400-Тестовый сценарий.docx"
    _build_docx(docx_path)

    proc = subprocess.run(
        [sys.executable, "tools/read_docx.py", str(docx_path)],
        text=True,
        capture_output=True,
        check=True,
    )
    payload = json.loads(proc.stdout)

    assert payload["title"] == "Заголовок ролика"
    assert payload["scheduled_at"] == "2026-06-25 14:00"
    assert payload["alias"] == "Тестовый сценарий"
    assert payload["body_text"] == "Первый абзац\n\nВторой абзац"


def test_read_docx_cli_missing_file_returns_json_error(tmp_path):
    proc = subprocess.run(
        [sys.executable, "tools/read_docx.py", str(tmp_path / "missing.docx")],
        text=True,
        capture_output=True,
    )

    assert proc.returncode == 1
    assert json.loads(proc.stderr)["error"].startswith("File not found:")
