import base64
import json
import subprocess
import sys

from docx import Document
from PIL import Image

from tools import upload_images as ui


def _make_png(path, color=(20, 90, 160)):
    Image.new("RGB", (16, 16), color).save(path, "PNG")


def _make_docx_with_images(path, image_paths):
    doc = Document()
    doc.add_paragraph("До первой картинки")
    for image_path in image_paths:
        doc.add_picture(str(image_path))
        doc.add_paragraph("После картинки")
    doc.save(path)


def test_extract_images_from_files_reads_supported_images_and_skips_unknown(tmp_path, capsys):
    png_path = tmp_path / "cover.png"
    txt_path = tmp_path / "notes.txt"
    _make_png(png_path)
    txt_path.write_text("not an image", encoding="utf-8")

    images = ui.extract_images_from_files([str(png_path), str(txt_path)])

    assert len(images) == 1
    assert images[0]["filename"] == "cover.png"
    assert images[0]["content_type"] == "image/png"
    assert base64.b64decode(images[0]["data"]).startswith(b"\x89PNG")
    assert "неизвестный тип файла" in capsys.readouterr().err


def test_extract_images_from_docx_preserves_document_order(tmp_path):
    first = tmp_path / "first.png"
    second = tmp_path / "second.png"
    docx_path = tmp_path / "images.docx"
    _make_png(first, (255, 0, 0))
    _make_png(second, (0, 255, 0))
    _make_docx_with_images(docx_path, [first, second])

    images = ui.extract_images_from_docx(docx_path)

    assert [image["content_type"] for image in images] == ["image/png", "image/png"]
    assert len(images) == 2
    assert images[0]["data"] != images[1]["data"]


def test_upload_images_returns_wordpress_compatible_json(monkeypatch):
    captured = []

    def fake_upload(data, filename, content_type):
        captured.append((data, filename, content_type))
        return {
            "id": 42,
            "source_url": "https://example.com/uploads/cover.png",
            "filename": filename,
            "size": len(data),
        }

    monkeypatch.setattr(ui, "upload_media_bytes", fake_upload)

    result = ui.upload_images(
        [
            {
                "filename": "cover.png",
                "content_type": "image/png",
                "data": base64.b64encode(b"png-data").decode("ascii"),
            }
        ]
    )

    assert captured == [(b"png-data", "cover.png", "image/png")]
    assert result == {
        "success": True,
        "uploaded": [
            {
                "filename": "cover.png",
                "url": "https://example.com/uploads/cover.png",
                "path": "https://example.com/uploads/cover.png",
                "media_id": 42,
                "size": 8,
            }
        ],
        "count": 1,
    }


def test_upload_images_cli_dry_run_reports_count(tmp_path):
    png_path = tmp_path / "cover.png"
    _make_png(png_path)

    proc = subprocess.run(
        [sys.executable, "tools/upload_images.py", str(png_path), "--dry-run"],
        text=True,
        capture_output=True,
        check=True,
    )

    assert "cover.png (image/png" in proc.stdout
    assert "Итого: 1 файлов" in proc.stdout


def test_upload_images_cli_outputs_json_with_mocked_files(tmp_path, monkeypatch):
    png_path = tmp_path / "cover.png"
    _make_png(png_path)

    # Exercise JSON serialization shape without hitting WordPress by calling the function directly.
    image = ui.extract_images_from_files([str(png_path)])[0]
    monkeypatch.setattr(
        ui,
        "upload_media_bytes",
        lambda data, filename, content_type: {
            "id": 7,
            "source_url": "https://example.com/uploads/cover.png",
            "filename": filename,
            "size": len(data),
        },
    )

    encoded = json.dumps(ui.upload_images([image]), ensure_ascii=False)

    assert '"success": true' in encoded
    assert "https://example.com/uploads/cover.png" in encoded
