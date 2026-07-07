import json
import subprocess
from pathlib import Path
from unittest.mock import patch

from docx import Document
from PIL import Image, ImageDraw

from tools import render_video_visuals as rv


def build_docx_with_embedded_images(path, tmp_path):
    first_image = tmp_path / "first.png"
    second_image = tmp_path / "second.png"
    Image.new("RGB", (640, 320), "#d64535").save(first_image)
    Image.new("RGB", (320, 640), "#2c62aa").save(second_image)

    doc = Document()
    doc.add_heading("ERP не спасает производство", level=1)
    doc.add_paragraph("Перед первой схемой нужно объяснить карту сжатия.")
    doc.add_picture(str(first_image))
    doc.add_paragraph("Схема 1. Карта сжатия показывает поток заказа.")
    doc.add_paragraph("Между визуалами идет модельный кейс.")
    doc.add_picture(str(second_image))
    doc.add_paragraph("График 1. Деньги возвращаются от сжатия потока.")
    doc.save(path)
    return path


def test_find_matching_docx_uses_video_basename(tmp_path):
    teleprompter = tmp_path / "teleprompter"
    teleprompter.mkdir()
    docx = teleprompter / "2026-05-01-1400-burovye-krs.docx"
    docx.write_bytes(b"placeholder")
    video = tmp_path / "2026-05-01-1400-burovye-krs.mp4"
    video.write_bytes(b"video")

    assert rv.find_matching_docx(video, teleprompter) == docx


def test_generate_visual_plan_creates_bounded_cards_with_source_anchors():
    body = "\n\n".join(
        [
            "Главная ошибка здесь не в рекламе, а в том, что система работает без ограничения.",
            "Сначала нужно найти бутылочное горлышко и посчитать деньги, которые там застряли.",
            "Потом ты переносишь внимание команды на один участок, а не размазываешь бюджет по всем отделам.",
            "Через 30 дней видно, где процесс ускорился и где появились новые потери.",
            "Финальный вывод простой: управляй ограничением, а не героизмом сотрудников.",
            "Если это игнорировать, бизнес каждый месяц теряет прибыль на ровном месте.",
        ]
    )

    plan = rv.generate_visual_plan("Тестовый ролик", body)

    assert 6 <= len(plan["slides"]) <= 10
    assert plan["slides"][0]["anchor"] in body
    assert all(slide["duration_sec"] >= 12 for slide in plan["slides"])
    assert all(1 <= len(slide["bullets"]) <= 4 for slide in plan["slides"])


def test_generate_visual_plan_splits_long_transcript_paragraph_into_cards():
    body = " ".join(
        f"Предложение {index} про ERP и производство показывает ограничение процесса."
        for index in range(1, 31)
    )

    plan = rv.generate_visual_plan("ERP", body)

    assert len(plan["slides"]) >= 6
    assert plan["slides"][0]["anchor"].startswith("Предложение 1")
    assert plan["slides"][1]["anchor"].startswith("Предложение")


def test_match_timeline_prefers_word_timestamps_and_falls_back():
    slides = [
        {
            "id": "slide_001",
            "anchor": "найти бутылочное горлышко",
            "duration_sec": 20,
        },
        {
            "id": "slide_002",
            "anchor": "этой фразы нет в транскрипте",
            "duration_sec": 20,
        },
    ]
    body_text = "Сначала нужно найти бутылочное горлышко и посчитать деньги. Потом другой блок."
    transcript = {
        "words": [
            {"text": "Сначала", "start": 0, "end": 400},
            {"text": "нужно", "start": 420, "end": 650},
            {"text": "найти", "start": 700, "end": 900},
            {"text": "бутылочное", "start": 920, "end": 1300},
            {"text": "горлышко", "start": 1320, "end": 1700},
        ]
    }

    timeline = rv.build_timeline(slides, transcript, body_text, video_duration=120)

    assert timeline[0]["match_method"] == "asr"
    assert timeline[0]["start_sec"] == 0.7
    assert timeline[1]["match_method"] == "fallback"
    assert timeline[1]["start_sec"] > timeline[0]["start_sec"]


def test_match_timeline_tries_anchor_candidates_before_fallback():
    slides = [
        {
            "id": "docx_visual_001",
            "anchor": "Схема 1. Этой подписи в речи нет",
            "anchor_candidates": [
                "Схема 1. Этой подписи в речи нет",
                "карта сжатия показывает поток заказа",
            ],
            "duration_sec": 20,
        }
    ]
    transcript = {
        "words": [
            {"text": "карта", "start": 2000, "end": 2300},
            {"text": "сжатия", "start": 2320, "end": 2600},
            {"text": "показывает", "start": 2620, "end": 2900},
            {"text": "поток", "start": 2920, "end": 3200},
            {"text": "заказа", "start": 3220, "end": 3500},
        ]
    }

    timeline = rv.build_timeline(slides, transcript, "карта сжатия показывает поток заказа", video_duration=80)

    assert timeline[0]["match_method"] == "asr"
    assert timeline[0]["start_sec"] == 2.0


def test_timeline_until_next_extends_visuals_to_next_start():
    slides = [
        {"id": "slide_001", "anchor": "первый визуал", "duration_sec": 10},
        {"id": "slide_002", "anchor": "второй визуал", "duration_sec": 10},
        {"id": "slide_003", "anchor": "третий визуал", "duration_sec": 10},
    ]
    transcript = {
        "words": [
            {"text": "первый", "start": 1000, "end": 1200},
            {"text": "визуал", "start": 1220, "end": 1500},
            {"text": "второй", "start": 21000, "end": 21200},
            {"text": "визуал", "start": 21220, "end": 21500},
            {"text": "третий", "start": 50000, "end": 50200},
            {"text": "визуал", "start": 50220, "end": 50500},
        ]
    }

    timeline = rv.build_timeline(
        slides,
        transcript,
        "первый визуал второй визуал третий визуал",
        video_duration=90,
        duration_mode="until-next",
    )

    assert timeline[0]["start_sec"] == 1.0
    assert timeline[0]["end_sec"] == 20.9
    assert timeline[1]["end_sec"] == 49.9
    assert timeline[2]["end_sec"] == 90


def test_timeline_until_next_keeps_docx_visuals_with_close_duplicate_anchors():
    timeline = rv.extend_timeline_until_next(
        [
            {"slide_id": "docx_visual_001", "start_sec": 10.0, "end_sec": 34.0},
            {"slide_id": "docx_visual_002", "start_sec": 15.1, "end_sec": 39.1},
            {"slide_id": "docx_visual_003", "start_sec": 50.0, "end_sec": 74.0},
        ],
        video_duration=100,
    )

    assert [item["slide_id"] for item in timeline] == [
        "docx_visual_001",
        "docx_visual_002",
        "docx_visual_003",
    ]
    assert timeline[0]["end_sec"] == 15.0


def test_build_timeline_until_next_enforces_five_second_minimum_for_close_visuals():
    slides = [
        {"id": "docx_visual_001", "anchor": "общий якорь", "duration_sec": 24},
        {"id": "docx_visual_002", "anchor": "общий якорь", "duration_sec": 24},
        {"id": "docx_visual_003", "anchor": "общий якорь", "duration_sec": 24},
    ]

    timeline = rv.build_timeline(
        slides,
        {"words": []},
        "общий якорь",
        video_duration=120,
        duration_mode="until-next",
    )

    assert [item["slide_id"] for item in timeline] == [
        "docx_visual_001",
        "docx_visual_002",
        "docx_visual_003",
    ]
    assert all(
        round(item["end_sec"] - item["start_sec"], 2) >= 5.0
        for item in timeline
    )


def test_render_slide_card_outputs_1920x1080_png(tmp_path):
    slide = {
        "id": "slide_001",
        "title": "Бутылочное горлышко",
        "kind": "principle",
        "bullets": ["Найди ограничение", "Считай деньги", "Не размазывай бюджет"],
    }

    output = rv.render_slide_card(slide, tmp_path)

    with Image.open(output) as image:
        assert image.size == (1920, 1080)
        assert image.mode == "RGB"


def test_render_video_with_fake_timeline_creates_output_mp4(tmp_path):
    raw = tmp_path / "2026-05-01-1400-test.mp4"
    subprocess.run(
        [
            "ffmpeg",
            "-y",
            "-f",
            "lavfi",
            "-i",
            "testsrc=size=1920x1080:rate=25",
            "-f",
            "lavfi",
            "-i",
            "sine=frequency=1000:sample_rate=44100",
            "-t",
            "2",
            "-pix_fmt",
            "yuv420p",
            "-c:v",
            "libx264",
            "-c:a",
            "aac",
            str(raw),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    slide = rv.render_slide_card(
        {
            "id": "slide_001",
            "title": "Тест",
            "kind": "metric",
            "bullets": ["Один тезис"],
        },
        tmp_path,
    )
    output = tmp_path / "out.mp4"

    rv.render_video(raw, [{"image_path": str(slide), "start_sec": 0.2, "end_sec": 1.2}], output)

    assert output.exists()
    assert output.stat().st_size > 0


def test_render_video_preserves_duration_when_plain_segments_are_between_keyframes(tmp_path):
    raw = tmp_path / "2026-05-01-1400-long-gop.mp4"
    subprocess.run(
        [
            "ffmpeg",
            "-y",
            "-f",
            "lavfi",
            "-i",
            "testsrc=size=640x360:rate=25",
            "-f",
            "lavfi",
            "-i",
            "sine=frequency=1000:sample_rate=44100",
            "-t",
            "8",
            "-pix_fmt",
            "yuv420p",
            "-c:v",
            "libx264",
            "-g",
            "250",
            "-sc_threshold",
            "0",
            "-c:a",
            "aac",
            str(raw),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    slide = rv.render_slide_card(
        {
            "id": "slide_001",
            "title": "Тест",
            "kind": "metric",
            "bullets": ["Один тезис"],
        },
        tmp_path,
    )
    output = tmp_path / "out.mp4"

    rv.render_video(raw, [{"image_path": str(slide), "start_sec": 2.0, "end_sec": 3.0}], output)

    assert abs(rv.ffprobe_duration(output) - rv.ffprobe_duration(raw)) < 0.25


def test_cli_dry_run_writes_artifacts_without_assemblyai(tmp_path):
    raw = tmp_path / "2026-05-01-1400-test.mp4"
    raw.write_bytes(b"not a real video for dry-run")
    docx_json = tmp_path / "docx.json"
    docx_json.write_text(
        json.dumps(
            {
                "title": "Тест",
                "body_text": "Первый тезис про деньги.\n\nВторой тезис про ограничение.\n\nТретий тезис про результат.",
                "scheduled_at": "2026-05-01 14:00",
                "alias": "test",
                "source_path": "fake.docx",
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    result = rv.run_pipeline(
        raw,
        docx_json_path=docx_json,
        output_dir=tmp_path / "rendered",
        work_root=tmp_path / "visuals",
        dry_run=True,
    )

    assert Path(result["slide_plan"]).exists()
    assert Path(result["timeline"]).exists()
    assert result["output_video"] is None


def test_dry_run_can_build_visuals_from_transcript_without_docx(tmp_path):
    raw = tmp_path / "2026-05-28 ЕРП Не спасает производство.mp4"
    raw.write_bytes(b"not a real video for dry-run")
    transcript_json = tmp_path / "transcript.json"
    transcript_json.write_text(
        json.dumps(
            {
                "text": "ЕРП не спасает производство, если процесс не описан. Сначала нужно найти ограничение. Потом убрать ручные обходы и проверить результат.",
                "words": [
                    {"text": "ЕРП", "start": 0, "end": 300},
                    {"text": "не", "start": 320, "end": 450},
                    {"text": "спасает", "start": 470, "end": 900},
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    result = rv.run_pipeline(
        raw,
        transcript_json_path=transcript_json,
        output_dir=tmp_path / "rendered",
        work_root=tmp_path / "visuals",
        dry_run=True,
    )

    slide_plan = json.loads(Path(result["slide_plan"]).read_text(encoding="utf-8"))
    assert slide_plan["title"] == "2026-05-28 ЕРП Не спасает производство"
    assert "ЕРП не спасает производство" in slide_plan["slides"][0]["anchor"]
    assert result["transcript"] == str(transcript_json)
    assert Path(result["timeline"]).exists()


def test_extract_docx_embedded_images_preserves_order_and_context(tmp_path):
    docx_path = build_docx_with_embedded_images(tmp_path / "visuals.docx", tmp_path)

    visuals = rv.extract_docx_embedded_images(docx_path, tmp_path / "extracted")

    assert [visual["id"] for visual in visuals] == ["docx_visual_001", "docx_visual_002"]
    assert visuals[0]["caption"] == "Схема 1. Карта сжатия показывает поток заказа."
    assert visuals[0]["paragraph_before"] == "Перед первой схемой нужно объяснить карту сжатия."
    assert visuals[0]["anchor"] == "Схема 1. Карта сжатия показывает поток заказа."
    assert visuals[1]["caption"] == "График 1. Деньги возвращаются от сжатия потока."
    assert Path(visuals[0]["source_image_path"]).exists()
    assert Path(visuals[1]["source_image_path"]).exists()


def test_render_docx_image_card_outputs_1920x1080_without_cropping(tmp_path):
    docx_path = build_docx_with_embedded_images(tmp_path / "visuals.docx", tmp_path)
    visual = rv.extract_docx_embedded_images(docx_path, tmp_path / "extracted")[0]

    output = rv.render_docx_image_card(visual, tmp_path / "cards")

    with Image.open(output) as image:
        assert image.size == (1920, 1080)
        assert image.mode == "RGB"


def test_render_docx_image_card_uses_video_title_footer_without_transcript_caption(tmp_path):
    docx_path = build_docx_with_embedded_images(tmp_path / "visuals.docx", tmp_path)
    visual = rv.extract_docx_embedded_images(docx_path, tmp_path / "extracted")[0]
    drawn_text = []
    original_text = ImageDraw.ImageDraw.text

    def capture_text(draw, xy, text, *args, **kwargs):
        drawn_text.append(text)
        return original_text(draw, xy, text, *args, **kwargs)

    with patch.object(ImageDraw.ImageDraw, "text", capture_text):
        rv.render_docx_image_card(visual, tmp_path / "cards", video_title="Миксовый сдвиг FMCG")

    assert visual["caption"] not in drawn_text
    assert "Сергей Веселков | Миксовый сдвиг FMCG" in drawn_text


def test_docx_visual_plan_can_use_short_footer_title_without_changing_plan_title(tmp_path):
    docx_path = build_docx_with_embedded_images(tmp_path / "visuals.docx", tmp_path)
    footer_titles = []

    def fake_render_card(visual, output_dir, video_title=None):
        footer_titles.append(video_title)
        output = Path(output_dir) / f"{visual['id']}.png"
        output.parent.mkdir(parents=True, exist_ok=True)
        Image.new("RGB", (1920, 1080), "#ffffff").save(output)
        return output

    with patch.object(rv, "render_docx_image_card", fake_render_card):
        plan = rv.generate_docx_image_visual_plan(
            "Полный заголовок ролика для YouTube",
            docx_path,
            tmp_path / "cards",
            footer_title="Миксовый сдвиг FMCG",
        )

    assert plan["title"] == "Полный заголовок ролика для YouTube"
    assert footer_titles == ["Миксовый сдвиг FMCG", "Миксовый сдвиг FMCG"]


def test_dry_run_can_build_visuals_from_docx_embedded_images(tmp_path):
    raw = tmp_path / "2026-06-11 1400 erp-ne-spasaet-proizvodstvo-blog.mp4"
    raw.write_bytes(b"not a real video for dry-run")
    docx_path = build_docx_with_embedded_images(tmp_path / "visuals.docx", tmp_path)
    transcript_json = tmp_path / "transcript.json"
    transcript_json.write_text(
        json.dumps(
            {
                "text": "Карта сжатия показывает поток заказа. Деньги возвращаются от сжатия потока.",
                "words": [
                    {"text": "Карта", "start": 1000, "end": 1300},
                    {"text": "сжатия", "start": 1320, "end": 1600},
                    {"text": "показывает", "start": 1620, "end": 1900},
                    {"text": "поток", "start": 1920, "end": 2200},
                    {"text": "заказа", "start": 2220, "end": 2500},
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    result = rv.run_pipeline(
        raw,
        docx_path=docx_path,
        transcript_json_path=transcript_json,
        output_dir=tmp_path / "rendered",
        work_root=tmp_path / "visuals",
        dry_run=True,
        use_docx_images=True,
    )

    visual_plan = json.loads(Path(result["slide_plan"]).read_text(encoding="utf-8"))
    assert len(visual_plan["slides"]) == 2
    assert visual_plan["slides"][0]["source"] == "docx_image"
    assert Path(visual_plan["slides"][0]["image_path"]).exists()
    assert result["asr_matches"] == 1


def test_placeholder_assemblyai_key_is_treated_as_missing():
    assert rv.is_missing_api_key("")
    assert rv.is_missing_api_key("your_assemblyai_api_key_here")
    assert not rv.is_missing_api_key("realistic_non_placeholder_key")


def test_render_with_visuals_contract_is_documented():
    command = Path(".codex/commands/render-with-visuals.md")
    skill = Path(".codex/skills/render-with-visuals/SKILL.md")
    env_example = Path(".env.example").read_text(encoding="utf-8")
    readme = Path("README.md").read_text(encoding="utf-8")
    agents = Path("AGENTS.md").read_text(encoding="utf-8")

    assert command.exists()
    assert skill.exists()
    command_text = command.read_text(encoding="utf-8")
    skill_text = skill.read_text(encoding="utf-8")
    assert "ASSEMBLYAI_API_KEY" in env_example
    assert "render-with-visuals" in readme
    assert "YYYY-MM-DD-HHMM-Русский заголовок.mp4" in readme
    assert "YYYY-MM-DD-HHMM-Русский заголовок.mp4" in command_text
    assert "YYYY-MM-DD-HHMM-Русский заголовок.mp4" in skill_text
    assert "YYYY-MM-DD-HHMM-alias.mp4" not in command_text
    assert "YYYY-MM-DD-HHMM-alias.mp4" not in skill_text
    assert "не raw" in agents


def test_render_docx_visuals_contract_is_documented():
    command = Path(".codex/commands/render-docx-visuals.md")
    skill = Path(".codex/skills/render-docx-visuals/SKILL.md")

    assert command.exists()
    assert skill.exists()

    command_text = command.read_text(encoding="utf-8")
    skill_text = skill.read_text(encoding="utf-8")
    assert "render-docx-visuals" in command_text
    assert "--use-docx-images" in skill_text
    assert "embedded_visual_plan.json" in skill_text
