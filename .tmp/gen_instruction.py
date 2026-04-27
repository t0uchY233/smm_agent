import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def add_step(number, emoji, title):
    para = doc.add_paragraph()
    run = para.add_run(f"{emoji}  ШАГ {number}. {title}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x9E)
    para.paragraph_format.space_before = Pt(20)
    para.paragraph_format.space_after = Pt(6)

def add_text(text, bold=False, size=12, gray=False, italic=False, indent=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if gray:
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.3
    if indent:
        para.paragraph_format.left_indent = Cm(0.8)

def add_bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    for run in para.runs:
        run.clear()
    para.clear()
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(2)

def add_sep():
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("_" * 50)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(10)

# ========= TITLE =========
h = doc.add_heading("Контент-конвейер: от идеи до публикации", level=0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Пошаговая инструкция")
r.font.name = "Arial"
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
add_text("Весь процесс \u2014 4 шага. Большая часть работы автоматическая.", bold=True, size=13)
add_text("Тебе нужно: придумать тему, записать видео и два раза нажать OK.", size=12)

doc.add_paragraph()
schema = doc.add_paragraph()
schema.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = schema.add_run("Тема --> Сценарий --> Запись --> YouTube --> Telegram + Блог + Дзен")
r.font.name = "Consolas"
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x1A, 0x5C, 0x9E)
r.bold = True

schema2 = doc.add_paragraph()
schema2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = schema2.add_run("  ты        ты + ИИ      ты       автомат      ты + ИИ")
r2.font.name = "Consolas"
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_sep()

# ========= STEP 0 =========
add_step(0, "\U0001F4DD", "Сценарий")
add_text("Генерируем сценарий и текст для телесуфлёра.", size=11)
add_text("Где: Claude Code (компьютер или телефон через claude.ai/code)", size=11, gray=True, italic=True)
add_text("")
add_text("Твои действия:", bold=True, size=11)
add_bullet("Открой Claude Code")
add_bullet('Напиши:  /video-script "тема твоего видео"')
add_bullet("Ответь на 2-3 вопроса (или пропусти \u2014 ИИ додумает сам)")
add_bullet("Прочитай сценарий. Нажми OK или попроси изменить")
add_bullet("Получи текст для телесуфлёра (в чате + файл .docx)")
add_text("")
add_text("Файл с телесуфлёром автоматически сохраняется в папку .tmp/teleprompter/", size=10, gray=True, italic=True)

add_sep()

# ========= STEP 1 =========
add_step(1, "\U0001F3AC", "Запись видео")
add_text("Записываем видео, читая текст с телесуфлёра.", size=11)
add_text("Где: камера или телефон", size=11, gray=True, italic=True)
add_text("")
add_text("Твои действия:", bold=True, size=11)
add_bullet("Открой файл .docx с текстом телесуфлёра")
add_bullet("Загрузи текст в приложение-суфлёр на телефоне")
add_bullet("Запиши видео, читая текст с экрана")
add_bullet("Сохрани видеофайл на Google Диск (в нужную папку)")
add_text("")
add_text('Имя файла: "Заголовок видео.mp4"', size=10, gray=True, italic=True)

add_sep()

# ========= STEP 2 =========
add_step(2, "\u2699\uFE0F", "Автоматическая обработка (ничего делать не надо)")
add_text("Робот (n8n) делает всё сам. Ты просто ждёшь.", size=11)
add_text("Где: сервер, полностью автоматически", size=11, gray=True, italic=True)
add_text("")
add_text("Что происходит без тебя:", bold=True, size=11)
add_bullet("Робот замечает новое видео на Google Диске")
add_bullet("Загружает его на YouTube (пока приватное)")
add_bullet("Генерирует обложку для видео (ИИ)")
add_bullet("Устанавливает обложку на YouTube")
add_bullet("Присылает тебе уведомление в Telegram со ссылкой")
add_text("")
add_text("Твои действия:", bold=True, size=11)
add_bullet("Подожди уведомления в Telegram (2-5 минут)")
add_bullet("Скопируй ссылку на YouTube из уведомления")

add_sep()

# ========= STEP 3 =========
add_step(3, "\U0001F4E2", "Публикация на все площадки")
add_text("Из видео создаём посты для Telegram, блога и Дзена.", size=11)
add_text("Где: Claude Code", size=11, gray=True, italic=True)
add_text("")
add_text("Твои действия:", bold=True, size=11)
add_bullet("Открой Claude Code")
add_bullet("Вставь ссылку на YouTube из уведомления")
add_bullet("ИИ сам вытащит текст из видео (субтитры)")
add_bullet("ИИ напишет пост для Telegram + статью для блога")
add_bullet("Прочитай тексты. Нажми OK или попроси исправить")
add_bullet("ИИ опубликует:")
add_text("     \u2192  Telegram-канал: пост с картинкой", size=11, indent=True)
add_text("     \u2192  Блог veselkov.me: полная статья", size=11, indent=True)
add_text("     \u2192  Яндекс Дзен: подхватит из блога автоматически", size=11, indent=True)

add_sep()

# ========= CHEAT SHEET TABLE =========
h2 = doc.add_heading("Шпаргалка", level=1)

table = doc.add_table(rows=5, cols=4)
table.style = "Medium Grid 1 Accent 1"

headers = ["Шаг", "Что делать", "Где", "Время"]
for i, ht in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ht
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.name = "Arial"
            r.bold = True

data = [
    ["0. Сценарий",   "Написать тему",          "Claude Code", "5 мин"],
    ["1. Съёмка",     "Записать видео",         "Камера",      "10-15 мин"],
    ["2. Обработка",  "Ничего (автомат)",        "n8n сервер",  "2-5 мин"],
    ["3. Публикация", "Вставить URL, нажать OK", "Claude Code", "3 мин"],
]

for ri, rd in enumerate(data):
    for ci, val in enumerate(rd):
        cell = table.rows[ri + 1].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                r.font.name = "Arial"

doc.add_paragraph()

# ========= COMMANDS =========
h3 = doc.add_heading("Команды Claude Code", level=1)

add_text('/video-script "тема"', bold=True, size=12)
add_text("Создать сценарий и телесуфлёр-текст", size=11, indent=True)
add_text("")
add_text("/publish-from-script", bold=True, size=12)
add_text("Запускается когда вставляешь YouTube URL", size=11, indent=True)

doc.add_paragraph()
add_sep()

# ========= FOOTER =========
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Итого твоих действий: тема + запись + два раза OK")
r.font.name = "Arial"
r.font.size = Pt(14)
r.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x5C, 0x9E)

foot2 = doc.add_paragraph()
foot2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = foot2.add_run("Всё остальное делают роботы.")
r2.font.name = "Arial"
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

output = ".tmp/teleprompter/instrukciya-content-pipeline.docx"
doc.save(output)
print(f"OK: {output}")
